# -*- coding: utf-8 -*-

import os
import re

from docx import Document


class DocxRenderer(object):
    TAG_RE = re.compile(r"\{[^{}]+\}", re.UNICODE)

    # Недопустимые XML-символы. Иногда они могут прийти из БД
    # и из-за них Word показывает сообщение о восстановлении документа.
    INVALID_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", re.UNICODE)

    def __init__(self, context_builder, tag_resolver):
        self.context_builder = context_builder
        self.tag_resolver = tag_resolver

        self.last_template_tags = []
        self.last_unknown_tags = []
        self.last_unresolved_tags = []

    def render_to_file(self, template_path, output_path, state):
        if not os.path.exists(template_path):
            raise FileNotFoundError(
                "Не найден шаблон документа: {0}".format(template_path)
            )

        context = self.context_builder.build(state)

        document = Document(template_path)

        self.last_template_tags = self._collect_tags_from_document(document)
        self.last_unknown_tags = self._collect_unknown_tags(document, context)

        self._process_document(document, context)

        self.last_unresolved_tags = self._collect_tags_from_document(document)

        output_dir = os.path.dirname(os.path.abspath(output_path))
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        document.save(output_path)

        return output_path

    def get_template_tags(self, template_path):
        if not os.path.exists(template_path):
            raise FileNotFoundError(
                "Не найден шаблон документа: {0}".format(template_path)
            )

        document = Document(template_path)
        return self._collect_tags_from_document(document)

    def get_unknown_tags_in_template(self, template_path, state):
        if not os.path.exists(template_path):
            raise FileNotFoundError(
                "Не найден шаблон документа: {0}".format(template_path)
            )

        context = self.context_builder.build(state)
        document = Document(template_path)

        return self._collect_unknown_tags(document, context)

    def _process_document(self, document, context):
        for paragraph in document.paragraphs:
            self._process_paragraph(paragraph, context)

        for table in document.tables:
            self._process_table(table, context)

        for section in document.sections:
            self._process_header_footer(section.header, context)
            self._process_header_footer(section.footer, context)

    def _process_header_footer(self, container, context):
        for paragraph in container.paragraphs:
            self._process_paragraph(paragraph, context)

        for table in container.tables:
            self._process_table(table, context)

    def _process_table(self, table, context):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, context)

                for nested_table in cell.tables:
                    self._process_table(nested_table, context)

    def _process_paragraph(self, paragraph, context):
        original_text = self._get_paragraph_text(paragraph)

        if not original_text:
            return

        if "{" not in original_text or "}" not in original_text:
            return

        if self._is_custom_intro_placeholder(original_text, context):
            rich_segments = context.get("__custom_intro_segments") or []

            if rich_segments:
                segments = self.tag_resolver.render_rich_segments(
                    rich_segments,
                    context,
                )
                self._set_paragraph_segments(paragraph, segments)
                return

        rendered_text = self.tag_resolver.render(original_text, context)
        rendered_text = self._clean_xml_text(rendered_text)

        if rendered_text == original_text:
            return

        if not rendered_text.strip() and self._paragraph_is_only_tags(original_text):
            self._set_paragraph_text(paragraph, "")
            return

        segments = self.tag_resolver.render_segments(original_text, context)
        self._set_paragraph_segments(paragraph, segments)

    @staticmethod
    def _is_custom_intro_placeholder(text, context):
        if not context.get("__use_custom_intro"):
            return False

        return (text or "").strip().lower() == "{вводная часть}"

    def _set_paragraph_segments(self, paragraph, segments):
        if not segments:
            self._set_paragraph_text(paragraph, "")
            return

        base_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")

        for run in paragraph.runs:
            run.text = ""

        first = True

        for segment in segments:
            text = self._clean_xml_text(segment.get("text") or "")
            if not text:
                continue

            is_bold = bool(segment.get("bold"))

            if first:
                run = base_run
                first = False
            else:
                run = paragraph.add_run()
                try:
                    run.style = base_run.style
                except Exception:
                    pass
                # Копируем шрифт из базового run-а (размер, гарнитура)
                if base_run.font.size:
                    run.font.size = base_run.font.size
                if base_run.font.name:
                    run.font.name = base_run.font.name

            run.text = text
            run.bold = True if is_bold else False

    @staticmethod
    def _get_paragraph_text(paragraph):
        if paragraph.runs:
            return "".join(run.text for run in paragraph.runs)

        return paragraph.text or ""

    def _set_paragraph_text(self, paragraph, text):
        text = self._clean_xml_text(text or "")

        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def _paragraph_is_only_tags(self, text):
        without_tags = self.TAG_RE.sub("", text or "")
        without_tags = without_tags.strip()
        without_tags = without_tags.strip(".,;:—-– ")

        return without_tags == ""

    def _collect_tags_from_document(self, document):
        tags = set()

        for paragraph in document.paragraphs:
            tags.update(self._collect_tags_from_text(self._get_paragraph_text(paragraph)))

        for table in document.tables:
            tags.update(self._collect_tags_from_table(table))

        for section in document.sections:
            tags.update(self._collect_tags_from_header_footer(section.header))
            tags.update(self._collect_tags_from_header_footer(section.footer))

        return sorted(tags)

    def _collect_tags_from_header_footer(self, container):
        tags = set()

        for paragraph in container.paragraphs:
            tags.update(self._collect_tags_from_text(self._get_paragraph_text(paragraph)))

        for table in container.tables:
            tags.update(self._collect_tags_from_table(table))

        return tags

    def _collect_tags_from_table(self, table):
        tags = set()

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    tags.update(self._collect_tags_from_text(self._get_paragraph_text(paragraph)))

                for nested_table in cell.tables:
                    tags.update(self._collect_tags_from_table(nested_table))

        return tags

    def _collect_tags_from_text(self, text):
        if not text:
            return []

        return self.TAG_RE.findall(text)

    def _collect_unknown_tags(self, document, context):
        tags = self._collect_tags_from_document(document)

        unknown = []
        for tag in tags:
            raw_tag = tag[1:-1].strip()
            if not self.tag_resolver.is_known_tag(raw_tag, context):
                unknown.append(tag)

        return sorted(set(unknown))

    def _clean_xml_text(self, text):
        return self.INVALID_XML_CHARS_RE.sub("", text or "")